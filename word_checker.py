import csv

from docx import Document
from string import punctuation


def get_common_words():

    common_words = set()

    with open('common_words.csv', 'r') as common_words_file:

        reader = csv.reader(common_words_file)

        for row in reader:
            row = [element.lower() for element in row]
            common_words.update(row)

        return common_words


def get_test_words(filename):

    test_words = {}
    word_doc = Document(filename)

    num_paragraphs = len(word_doc.paragraphs)

    print(word_doc.core_properties.identifier)

    for paragraph in range(0, num_paragraphs):
        text = word_doc.paragraphs[paragraph].text

        words = text.translate(str.maketrans('', '', punctuation)).split()

        words = [word.lower() for word in words]

        for word in words:

            if word in test_words.keys():

                updated_count = test_words[word] + 1
                test_words.update({word: updated_count})
            else:
                test_words.update({word: 1})

    return test_words


def verify(filename, test_words, common_words):

    rare_words = {}

    for word in test_words:

        if word not in common_words:
            rare_words.update({word: test_words[word]})

    if not bool(rare_words):
        print("Your file only contains common words.")
    else:
        print("Here are the rare words in your file:")

        for word in rare_words:

            print(f"Word: {word} \t | Occurences: {rare_words[word]}")


def get_filename():

    prompt = "Please enter the file name: "
    filename = input(prompt)

    if filename.strip() == '':
        print("File name is empty. Using default test file.")
        filename = "test.docx"
    else:
        print(f"You entered the following file name: {filename}")

    return filename


filename = get_filename()
test_words = get_test_words(filename)
common_words = get_common_words()
verify(filename, test_words, common_words)
